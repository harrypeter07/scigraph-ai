"""Script to generate a comprehensive Microsoft Word (.docx) for the Major Project Seminar Dossier."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls


def create_seminar_word_doc(output_path: str = "docs/SciGraph_Major_Project_Seminar_Dossier.docx"):
    doc = Document()

    # Page Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title
    t = doc.add_heading("SciGraph AI — 7th Semester Major Project Dossier", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(14, 116, 144)

    sub = doc.add_paragraph("Leakage-Free Scientific Impact Trajectory Forecasting via Heterogeneous Graph Neural Networks")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # 1. Project Aim & Problem Statement
    h1 = doc.add_heading("1. Project Aim & Problem Statement", level=1)
    for r in h1.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    p_aim = doc.add_paragraph(
        "• Project Aim: To design, implement, and empirically validate an end-to-end, zero-leakage Heterogeneous Graph Neural Network framework capable of forecasting the 5-year long-term citation impact trajectory of scientific publications at the moment of release, and to provide an interactive retrospective verification engine that benchmarks predictions against ground-truth historical outcomes.\n\n"
        "• Problem Statement: Existing bibliometric systems (Google Scholar, Semantic Scholar) are purely retrospective—they only show historical citation accumulations and cannot forecast trajectory on Day 1. Furthermore, existing machine learning literature frequently exhibits temporal data leakage (accidentally using future graph edges to inflate reported accuracy). Finally, raw citation counts vary drastically across fields and eras, rendering static thresholds invalid."
    )
    p_aim.runs[0].font.size = Pt(10.5)

    # 2. Measurable Objectives
    h2 = doc.add_heading("2. Measurable Technical Objectives", level=1)
    for r in h2.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    objs = [
        ("Zero-Leakage Ingestion", "Ingest multi-relational academic bibliographies from OpenAlex and enforce a strict observation cutoff (T_cutoff = Y_pub) that masks all future citations."),
        ("Dynamic Cohort Labeling", "Discretize 5-year citation growth into 3 cohort percentiles: Low (<50th), Medium (50th-90th), and High (>=90th landmark papers)."),
        ("Heterogeneous Graph Construction", "Build PyTorch Geometric HeteroData graphs connecting Papers, Authors, Institutions, and Topics across 4 relational edge types."),
        ("GNN Architecture Design", "Implement HeteroGraphSAGE and HeteroGAT in PyTorch with 5 non-leaking topological features."),
        ("Baseline-Anchored Evaluation", "Structurally pair all model metrics with a Majority-Class Baseline on identical temporal splits."),
        ("Retrospective Verification", "Engineer a queryable simulation engine that reconstructs snapshots at T_cutoff and benchmarks forecasts side-by-side with genuine outcomes.")
    ]
    for name, desc in objs:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {name}: ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(14, 116, 144)
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10)

    # 3. Literature Review Matrix
    h3 = doc.add_heading("3. Literature Review & Research Gap Matrix", level=1)
    for r in h3.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    tbl_lit = doc.add_table(rows=1, cols=4)
    tbl_lit.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_lit.style = 'Table Grid'

    headers_lit = ["Approach / Literature", "Platform / Source", "Core Weakness / Gap", "SciGraph AI Differentiation"]
    for i, title_text in enumerate(headers_lit):
        c = tbl_lit.rows[0].cells[i]
        c.text = title_text
        c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="0E7490"/>'.format(nsdecls('w'))))
        for run in c.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)

    lit_data = [
        ("Commercial Bibliometrics (Google Scholar, Scopus)", "Commercial Platforms (2004-Present)", "Zero Forecasting: Completely backward-looking.", "Day-1 Forecasting: Predicts 5-year cohort percentiles before citations accumulate."),
        ("Tabular ML & Text Baselines (Random Forest, SciBERT)", "Traditional Bibliometrics (2018-2022)", "Ignores Relational Topology: Fails to capture multi-hop author and institutional networks.", "Heterogeneous GNN: Propagates messages across author, institution, and topic nodes."),
        ("AGSTA-NET & H2CGL", "IEEE / ACM Papers (2023-2024)", "Homogeneous Graph Assumption: Ignores institution nodes; no interactive product.", "Multi-Relational Graph: Distinguishes 4 entity types + interactive verification UI."),
        ("BA-Cite (Bias-Aware Citation Forecasting)", "KDD / arXiv (Oct 2025)", "Research Prototype Only: Static script without queryable retrospective simulation.", "Interactive Verification Engine: Query any real OpenAlex paper and verify live.")
    ]

    for row_idx, (c0, c1, c2, c3) in enumerate(lit_data):
        cells = tbl_lit.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = c0, c1, c2, c3
        if row_idx % 2 == 1:
            for cell in cells:
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w'))))
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # 4. Methodology & Formulas
    h4 = doc.add_heading("4. System Methodology & Mathematical Equations", level=1)
    for r in h4.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    p_math = doc.add_paragraph(
        "• 5-Dim Feature Vector x: [x0: Normalized Year, x1: Early Citations, x2: Title Complexity, x3: Author Team Ratio, x4: ln(1 + Citations_cutoff)]\n"
        "• Hidden Layer Forward Pass: h1 = ReLU(W1 * x + b1) where W1 is (64 x 5) and b1 is (64 x 1)\n"
        "• Classifier Logits: z = W2 * h1 + b2 = [z0, z1, z2]^T where W2 is (3 x 64)\n"
        "• Softmax Probabilities: P(Class i) = exp(zi) / sum(exp(zj)) across Low (<50%), Medium (50-90%), and High (>=90%)\n"
        "• Decision Rule: Predicted Class = argmax P(Class i)"
    )
    p_math.runs[0].font.size = Pt(10)

    # 5. Results & Benchmark Table
    h5 = doc.add_heading("5. Empirical Benchmark Results & Verification", level=1)
    for r in h5.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    tbl_res = doc.add_table(rows=1, cols=5)
    tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_res.style = 'Table Grid'

    headers_res = ["Model Architecture", "Test Split", "Accuracy Fraction", "Accuracy %", "Macro-F1"]
    for i, title_text in enumerate(headers_res):
        c = tbl_res.rows[0].cells[i]
        c.text = title_text
        c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="0E7490"/>'.format(nsdecls('w'))))
        for run in c.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)

    res_data = [
        ("Majority-Class Baseline", "Time-Consistent (Y >= 2020)", "3/5", "60.0%", "0.3750"),
        ("Logistic Regression", "Time-Consistent (Y >= 2020)", "3/5", "60.0%", "0.4286"),
        ("Gradient Boosting (GBDT)", "Time-Consistent (Y >= 2020)", "3/5", "60.0%", "0.4286"),
        ("HeteroGraphSAGE (PyTorch GNN)", "Time-Consistent (Y >= 2020)", "3/5", "60.0%", "0.4286"),
        ("Tier 4 Multimodal Ablation", "Full Graph Topology", "4/5", "80.0%", "0.6667")
    ]

    for row_idx, (c0, c1, c2, c3, c4) in enumerate(res_data):
        cells = tbl_res.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text, cells[4].text = c0, c1, c2, c3, c4
        if row_idx % 2 == 1:
            for cell in cells:
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w'))))
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)

    doc.add_paragraph()

    # 6. Future Scope & Roadmap
    h6 = doc.add_heading("6. Future Scope & Research Roadmap", level=1)
    for r in h6.runs: r.font.color.rgb = RGBColor(15, 23, 42)

    p_future = doc.add_paragraph(
        "1. GPU Dataset Scale-Up: Ingest 5,000+ works via OpenAlex on Google Colab CUDA GPU to achieve high statistical power.\n"
        "2. Multimodal Dense Semantic Fusion: Combine GraphSAGE topology with 768-dimensional SPECTER / SciBERT abstract embeddings.\n"
        "3. Dynamic Temporal Attention: Implement Time-Aware GATv2 to model aging decay of author influence over time.\n"
        "4. Research Publication: Target an IEEE / Scopus-indexed student conference on Graph Data Mining & Scientometrics."
    )
    p_future.runs[0].font.size = Pt(10.5)

    doc.save(output_path)
    print(f"Seminar Word document saved to: {output_path}")


if __name__ == "__main__":
    create_seminar_word_doc()
