"""Script to generate a Microsoft Word (.docx) and standalone HTML document for SciGraph AI Mathematical Specification."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def create_word_doc(output_path: str = "docs/SciGraph_Mathematical_Formulas_and_Flowchart.docx"):
    doc = Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_heading("SciGraph AI — Mathematical Formulas, Features & Architecture Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(14, 116, 144)  # Deep Cyan

    # Subtitle
    sub = doc.add_paragraph("Official Academic & Mathematical Reference Guide for Seminar Presentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # Section 1: 5-Step System Pipeline Flow
    h1 = doc.add_heading("1. End-to-End System Architecture Flow", level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)

    p_flow = doc.add_paragraph(
        "The SciGraph AI pipeline executes in 5 sequential stages:\n"
        "1. OpenAlex Ingestion: Retrieves genuine paper metadata (Publication Year, Authors, Citations, Institutions).\n"
        "2. Temporal Snapshot (T_cutoff = Y_pub): Strictly masks all future citations post-publication to prevent leakage.\n"
        "3. Feature Vector Extraction (x): Constructs a 5-dimensional normalized vector.\n"
        "4. PyTorch HeteroGraphSAGE Forward Pass: Computes 64-dim hidden representations and 3-class Softmax probabilities.\n"
        "5. Dynamic Cohort Classification: Maps output to Low (<50%), Medium (50-90%), or High (>=90%) percentile tiers."
    )
    p_flow.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()

    # Section 2: Deep Dive into the 5 Extracted Features
    h2 = doc.add_heading("2. Deep Dive: The 5 Extracted Features (Formulas & Rationale)", level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    headers = ["Feature", "Mathematical Formula", "Why We Extracted This?", "What It Tells the Model"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        shading = parse_xml(r'<w:shd {} w:fill="0E7490"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

    features_data = [
        (
            "x0: Normalized Year",
            "x0 = (PubYear - 2012) / 10",
            "Publishing velocity and citation patterns change over decades.",
            "Anchors the paper in its temporal era, scaling years [2012, 2022] into [0.0, 1.0]."
        ),
        (
            "x1: Scaled Early Citations",
            "x1 = Citations_at_cutoff / 100",
            "Early citation rate in the release year indicates immediate community interest.",
            "Measures early adoption speed. Dividing by 100 normalizes values for neural gradients."
        ),
        (
            "x2: Title Complexity Ratio",
            "x2 = Title_Word_Count / 20",
            "Scientometric studies show concise titles correlate with broad fundamental breakthroughs.",
            "Measures the lexical granularity and scope of the research contribution."
        ),
        (
            "x3: Co-Author Team Ratio",
            "x3 = CoAuthor_Count / 10",
            "Modern scientific breakthroughs rely on multi-lab collaboration networks.",
            "Indicates the collaborative breadth and institutional reach behind the paper."
        ),
        (
            "x4: Log-Scaled Citation Velocity",
            "x4 = ln(1 + Citations_at_cutoff)",
            "Citations follow a power-law distribution (0.1% receive thousands of citations).",
            "Compresses extreme citation outliers so viral papers do not explode model weights."
        )
    ]

    for row_idx, (f_name, f_form, f_why, f_what) in enumerate(features_data):
        row_cells = table.add_row().cells
        row_cells[0].text = f_name
        row_cells[1].text = f_form
        row_cells[2].text = f_why
        row_cells[3].text = f_what

        # Alternate row background
        if row_idx % 2 == 1:
            for cell in row_cells:
                shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)

        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)

    doc.add_paragraph()

    # Section 3: PyTorch Neural Network Feed-Forward Equations
    h3 = doc.add_heading("3. PyTorch Neural Network Feed-Forward Equations", level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "Our PyTorch GraphSAGE neural network (ml/gnn/models/graphsage.py) converts the 5-dimensional feature vector into a 3-class probability distribution across 4 exact mathematical steps:"
    )

    eq_steps = [
        ("Step 1: Input Feature Vector", "x = [x0, x1, x2, x3, x4]^T  (Dimension: 5 x 1)"),
        ("Step 2: Hidden Layer Linear Transformation", "h1 = W1 * x + b1\nWhere W1 is a (64 x 5) trained weight matrix and b1 is a (64 x 1) bias vector."),
        ("Step 3: Non-Linear Activation Function (ReLU)", "a1 = ReLU(h1) = max(0, h1)\nAllows the neural network to learn non-linear patterns in citation dynamics."),
        ("Step 4: Output Classifier Layer (Logits)", "z = W2 * a1 + b2 = [z0, z1, z2]^T\nWhere W2 is a (3 x 64) trained weight matrix and b2 is a (3 x 1) bias vector. z represents the raw unnormalized scores."),
        ("Step 5: Softmax Probability Distribution", "P(Low)    = exp(z0) / (exp(z0) + exp(z1) + exp(z2))\nP(Medium) = exp(z1) / (exp(z0) + exp(z1) + exp(z2))\nP(High)   = exp(z2) / (exp(z0) + exp(z1) + exp(z2))\nTotal Probability Sum = P(Low) + P(Medium) + P(High) = 1.0 (100%)"),
        ("Step 6: Final Decision Rule (Argmax)", "Predicted Class = argmax { P(Low), P(Medium), P(High) }\nThe model chooses the single class with the highest probability score.")
    ]

    for title_s, body_s in eq_steps:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title_s}:\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(14, 116, 144)
        r_b = p.add_run(f"   {body_s}")
        r_b.font.size = Pt(10)

    doc.add_paragraph()

    # Section 4: Dynamic Cohort Percentiles
    h4 = doc.add_heading("4. Dynamic Cohort-Normalized Impact Classes", level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)

    p_cohort = doc.add_paragraph(
        "Instead of unfair hardcoded citation counts, we dynamically calculate the 5-year citation gain (Delta Citations_5 = Citations(Y+5) - Citations(Y)) and find the percentile thresholds within that specific year's cohort:\n\n"
        "• Class 0 (Low Impact): Delta Citations_5 < 50th Percentile (Bottom 50% of papers in that year)\n"
        "• Class 1 (Medium Impact): 50th Percentile <= Delta Citations_5 < 90th Percentile (Mainstream solid contribution)\n"
        "• Class 2 (High Impact): Delta Citations_5 >= 90th Percentile (Top 10% Landmark Breakthrough Papers)"
    )
    p_cohort.runs[0].font.size = Pt(10.5)

    doc.save(output_path)
    print(f"Word document successfully saved to: {output_path}")


def create_standalone_html(output_path: str = "docs/SciGraph_Mathematical_Specification.html"):
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SciGraph AI — Mathematical Formulas & Architecture Flowchart</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&family=JetBrains+Mono:wght@400;600&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-color: #0b0f19;
            --card-bg: rgba(17, 24, 39, 0.85);
            --accent-cyan: #00f2fe;
            --accent-blue: #4facfe;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
            --border-color: rgba(255, 255, 255, 0.1);
        }
        body {
            font-family: 'Outfit', sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 2.5rem 1.5rem;
            line-height: 1.6;
        }
        .container {
            max-width: 960px;
            margin: 0 auto;
        }
        .header {
            text-align: center;
            margin-bottom: 2.5rem;
        }
        .header h1 {
            font-size: 2.3rem;
            background: linear-gradient(135deg, var(--accent-cyan), var(--accent-blue));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 0.5rem;
        }
        .card {
            background: var(--card-bg);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 2rem;
            margin-bottom: 2rem;
            backdrop-filter: blur(12px);
            box-shadow: 0 8px 32px rgba(0,0,0,0.3);
        }
        h2 {
            color: var(--accent-cyan);
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 0.6rem;
            margin-top: 0;
            font-size: 1.4rem;
        }
        .flow-step {
            display: flex;
            align-items: flex-start;
            margin-bottom: 1.2rem;
        }
        .flow-num {
            background: linear-gradient(135deg, #00f2fe, #4facfe);
            color: #0b0f19;
            font-weight: 700;
            width: 32px;
            height: 32px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 1rem;
            flex-shrink: 0;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 1rem;
            font-size: 0.95rem;
        }
        th, td {
            padding: 0.9rem 1rem;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        th {
            background: rgba(0, 242, 254, 0.1);
            color: var(--accent-cyan);
            font-weight: 600;
        }
        code, .formula-box {
            font-family: 'JetBrains Mono', monospace;
            background: rgba(0, 0, 0, 0.4);
            border: 1px solid rgba(255,255,255,0.08);
            border-radius: 8px;
            padding: 0.3rem 0.6rem;
            color: #38bdf8;
            font-size: 0.95rem;
        }
        .formula-block {
            font-family: 'JetBrains Mono', monospace;
            background: rgba(0, 0, 0, 0.5);
            border-left: 4px solid var(--accent-cyan);
            padding: 1.2rem;
            border-radius: 8px;
            margin: 1rem 0;
            color: #7dd3fc;
            font-size: 0.95rem;
            line-height: 1.7;
        }
        .badge {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: 600;
        }
        .badge-low { background: rgba(168, 85, 247, 0.2); color: #c084fc; border: 1px solid #c084fc; }
        .badge-med { background: rgba(59, 130, 246, 0.2); color: #60a5fa; border: 1px solid #60a5fa; }
        .badge-high { background: rgba(20, 184, 166, 0.2); color: #2dd4bf; border: 1px solid #2dd4bf; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>SciGraph AI — Mathematical Formulas & Architecture</h1>
            <p style="color: var(--text-muted);">Interactive, Clean Reference Document for Seminar Presentation</p>
        </div>

        <div class="card">
            <h2>1. End-to-End System Architecture Flow</h2>
            <div class="flow-step">
                <div class="flow-num">1</div>
                <div><strong>OpenAlex Bibliographic Ingestion:</strong> Retrieves genuine paper metadata (Publication Year, Authors, Citations, Institutions).</div>
            </div>
            <div class="flow-step">
                <div class="flow-num">2</div>
                <div><strong>Observation Freeze (T_cutoff = Y_pub):</strong> Slices historical arrays. All future citations post-publication are strictly masked to guarantee zero leakage.</div>
            </div>
            <div class="flow-step">
                <div class="flow-num">3</div>
                <div><strong>Feature Extraction (Vector x):</strong> Assembles a 5-dimensional normalized numerical vector.</div>
            </div>
            <div class="flow-step">
                <div class="flow-num">4</div>
                <div><strong>PyTorch GraphSAGE Forward Pass:</strong> Passes vector through a 64-neuron hidden layer + ReLU activation + 3-class Softmax layer.</div>
            </div>
            <div class="flow-step">
                <div class="flow-num">5</div>
                <div><strong>Dynamic Cohort Classification:</strong> Selects predicted tier (Low &lt;50%, Medium 50-90%, High &ge;90%).</div>
            </div>
        </div>

        <div class="card">
            <h2>2. Deep Dive: The 5 Extracted Features</h2>
            <table>
                <thead>
                    <tr>
                        <th>Feature</th>
                        <th>Exact Formula</th>
                        <th>Why We Extracted This?</th>
                        <th>What It Tells the Model</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>x0: Normalized Year</strong></td>
                        <td><code>x0 = (PubYear - 2012) / 10</code></td>
                        <td>Publishing velocity changes over decades.</td>
                        <td>Anchors paper in its temporal era ([2012-2022] &rarr; [0.0, 1.0]).</td>
                    </tr>
                    <tr>
                        <td><strong>x1: Scaled Early Citations</strong></td>
                        <td><code>x1 = Citations_cutoff / 100</code></td>
                        <td>Early pickup indicates immediate peer interest.</td>
                        <td>Measures early adoption velocity normalized for neural gradients.</td>
                    </tr>
                    <tr>
                        <td><strong>x2: Title Complexity</strong></td>
                        <td><code>x2 = Title_Word_Count / 20</code></td>
                        <td>Concise titles correlate with broad fundamental breakthroughs.</td>
                        <td>Measures lexical scope and granularity of the paper.</td>
                    </tr>
                    <tr>
                        <td><strong>x3: Author Team Ratio</strong></td>
                        <td><code>x3 = CoAuthor_Count / 10</code></td>
                        <td>Major breakthroughs are heavily collaborative.</td>
                        <td>Indicates collaborative breadth and institutional reach.</td>
                    </tr>
                    <tr>
                        <td><strong>x4: Log-Citation Velocity</strong></td>
                        <td><code>x4 = ln(1 + Citations_cutoff)</code></td>
                        <td>Citations follow a heavy-tailed power-law distribution.</td>
                        <td>Compresses extreme outliers so viral papers don't blow up weights.</td>
                    </tr>
                </tbody>
            </table>
        </div>

        <div class="card">
            <h2>3. PyTorch Neural Network Feed-Forward Equations</h2>
            <div class="formula-block">
                <strong>Step 1: Input Vector</strong><br>
                x = [x0, x1, x2, x3, x4]^T  &isin; &reals;<sup>5 &times; 1</sup><br><br>

                <strong>Step 2: Hidden Layer Linear Transformation (64 neurons)</strong><br>
                h1 = W1 &middot; x + b1    [where W1 &isin; &reals;<sup>64 &times; 5</sup>, b1 &isin; &reals;<sup>64</sup>]<br><br>

                <strong>Step 3: Non-Linear Activation (ReLU)</strong><br>
                a1 = max(0, h1) = ReLU(h1)<br><br>

                <strong>Step 4: Output Classifier Layer (Logits)</strong><br>
                z = W2 &middot; a1 + b2 = [z0, z1, z2]^T    [where W2 &isin; &reals;<sup>3 &times; 64</sup>, b2 &isin; &reals;<sup>3</sup>]<br><br>

                <strong>Step 5: Softmax Probability Distribution (Sums to 100%)</strong><br>
                P(Low Impact)    = exp(z0) / [exp(z0) + exp(z1) + exp(z2)]<br>
                P(Medium Impact) = exp(z1) / [exp(z0) + exp(z1) + exp(z2)]<br>
                P(High Impact)   = exp(z2) / [exp(z0) + exp(z1) + exp(z2)]<br><br>

                <strong>Step 6: Argmax Decision Rule</strong><br>
                Predicted Class = argmax { P(Low), P(Medium), P(High) }
            </div>
        </div>

        <div class="card">
            <h2>4. Dynamic Cohort-Normalized Impact Classes</h2>
            <p>We measure 5-year growth <code>&Delta;Citations<sub>5</sub> = Citations(Y+5) - Citations(Y)</code> and compare against that year's cohort:</p>
            <ul>
                <li style="margin-bottom: 0.8rem;"><span class="badge badge-low">Class 0: Low Impact</span> &mdash; <code>&Delta;Citations &lt; 50th Percentile</code> (Bottom half of papers that year)</li>
                <li style="margin-bottom: 0.8rem;"><span class="badge badge-med">Class 1: Medium Impact</span> &mdash; <code>50th &le; &Delta;Citations &lt; 90th Percentile</code> (Mainstream solid contribution)</li>
                <li style="margin-bottom: 0.8rem;"><span class="badge badge-high">Class 2: High Impact</span> &mdash; <code>&Delta;Citations &ge; 90th Percentile</code> (Top 10% Landmark Breakthroughs)</li>
            </ul>
        </div>
    </div>
</body>
</html>
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"HTML document successfully saved to: {output_path}")


if __name__ == "__main__":
    os.makedirs("docs", exist_ok=True)
    create_word_doc()
    create_standalone_html()
